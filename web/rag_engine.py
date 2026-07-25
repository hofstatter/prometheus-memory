#!/usr/bin/env python3
"""Prometheus RAG Engine — sqlite-vec + fastembed + langchain chunker + OCR."""
import os, re, uuid
from pathlib import Path
from datetime import datetime
from langchain_text_splitters import RecursiveCharacterTextSplitter
from fastembed import TextEmbedding

PDF_MAX_PAGES = 200

MNEMOSYNE_HOME = Path(os.environ.get("MNEMOSYNE_HOME", Path.home() / ".hermes" / "mnemosyne"))
MNEMOSYNE_DB = Path(os.environ.get("PROMETHEUS_DB", MNEMOSYNE_HOME / "data" / "mnemosyne.db"))
CHUNK_SIZE, CHUNK_OVERLAP = 500, 50
EMBED_DIM = 384

_engine = None

def get_engine():
    global _engine
    if _engine is None:
        _engine = RAGEngine(str(MNEMOSYNE_DB))
    return _engine

class RAGEngine:
    def __init__(self, db_path):
        self.db_path = db_path
        self.embedder = TextEmbedding(model_name="sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2")
        self.splitter = RecursiveCharacterTextSplitter(
            chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP,
            separators=["\n\n", "\n", ". ", " ", ""]
        )
        self._init_tables()

    def _db(self):
        import sqlite3, sqlite_vec
        conn = sqlite3.connect(self.db_path, timeout=10)
        conn.execute("PRAGMA journal_mode=WAL")
        conn.execute("PRAGMA busy_timeout=5000")
        conn.execute("PRAGMA synchronous=NORMAL")
        conn.enable_load_extension(True)
        sqlite_vec.load(conn)
        conn.enable_load_extension(False)
        return conn

    def _init_tables(self):
        db = self._db()
        db.executescript("""
            CREATE TABLE IF NOT EXISTS rag_collections (
                id TEXT PRIMARY KEY, name TEXT NOT NULL UNIQUE,
                description TEXT DEFAULT '', created_at TEXT DEFAULT (datetime('now'))
            );
            CREATE TABLE IF NOT EXISTS rag_documents (
                id TEXT PRIMARY KEY, collection_id TEXT NOT NULL REFERENCES rag_collections(id) ON DELETE CASCADE,
                filename TEXT NOT NULL, mime_type TEXT NOT NULL, char_count INTEGER DEFAULT 0,
                chunk_count INTEGER DEFAULT 0, content_preview TEXT DEFAULT '',
                filepath TEXT DEFAULT '', created_at TEXT DEFAULT (datetime('now'))
            );
            CREATE TABLE IF NOT EXISTS rag_chunks (
                id TEXT PRIMARY KEY, document_id TEXT NOT NULL,
                chunk_index INTEGER NOT NULL, content TEXT NOT NULL,
                embedding BLOB
            );
            CREATE INDEX IF NOT EXISTS idx_rag_docs_collection ON rag_documents(collection_id);
            CREATE INDEX IF NOT EXISTS idx_rag_chunks_doc ON rag_chunks(document_id);
        """)
        db.commit()
        db.close()

    def create_collection(self, name, description=""):
        cid = re.sub(r"[^a-z0-9-]", "-", name.lower())[:40]
        db = self._db()
        db.execute("INSERT OR IGNORE INTO rag_collections(id,name,description) VALUES(?,?,?)",
                   (cid, name, description))
        db.commit()
        db.close()
        return cid

    def list_collections(self):
        db = self._db()
        rows = db.execute("SELECT id,name,description,created_at FROM rag_collections ORDER BY created_at DESC").fetchall()
        db.close()
        return [{"id": r[0], "name": r[1], "description": r[2], "created_at": r[3]} for r in rows]

    def index_text(self, collection_id, filename, text, filepath=""):
        doc_id = str(uuid.uuid4())[:12]
        chunks = self.splitter.split_text(text)
        if not chunks:
            return None
        embeddings = list(self.embedder.embed(chunks))
        db = self._db()
        db.execute(
            "INSERT INTO rag_documents(id,collection_id,filename,mime_type,char_count,chunk_count,content_preview,filepath) VALUES(?,?,?,?,?,?,?,?)",
            (doc_id, collection_id, filename, "text/plain", len(text), len(chunks), text[:200], filepath)
        )
        for i, (chunk, emb) in enumerate(zip(chunks, embeddings)):
            cid = f"{doc_id}_{i}"
            emb_blob = emb.tobytes()
            db.execute("INSERT INTO rag_chunks(id,document_id,chunk_index,content,embedding) VALUES(?,?,?,?,?)",
                       (cid, doc_id, i, chunk, emb_blob))
        db.commit()
        db.close()
        return {"document_id": doc_id, "chunks": len(chunks), "chars": len(text)}

    def extract_file(self, filepath):
        ext = Path(filepath).suffix.lower()
        if ext == ".pdf":
            import pypdfium2 as pdfium
            texts = []
            pdf = pdfium.PdfDocument(filepath)
            try:
                for i in range(min(len(pdf), PDF_MAX_PAGES)):
                    texts.append(pdf.get_page(i).get_textpage().get_text_range())
            finally:
                pdf.close()
            text = "\n".join(texts)
            if len(text.strip()) < 100:
                text = self._ocr_pdf(filepath)
            return text
        elif ext in (".png", ".jpg", ".jpeg"):
            import subprocess
            r = subprocess.run(["tesseract", filepath, "stdout", "-l", "por+eng"],
                               capture_output=True, text=True, timeout=60)
            return r.stdout.strip()
        elif ext == ".docx":
            from docx import Document
            return "\n".join(p.text for p in Document(filepath).paragraphs)
        else:
            return Path(filepath).read_text(errors="ignore")

    def _ocr_pdf(self, filepath):
        """OCR em PDF escaneado: renderiza páginas via pypdfium2 e roda tesseract por página."""
        import subprocess, tempfile
        import pypdfium2 as pdfium
        texts = []
        try:
            pdf = pdfium.PdfDocument(filepath)
            with tempfile.TemporaryDirectory() as tmp:
                for i in range(min(len(pdf), 30)):
                    bitmap = pdf.get_page(i).render(scale=200 / 72)
                    img_path = Path(tmp) / f"page_{i}.png"
                    bitmap.to_pil().save(str(img_path))
                    r = subprocess.run(["tesseract", str(img_path), "stdout", "-l", "por+eng"],
                                       capture_output=True, text=True, timeout=60)
                    texts.append(r.stdout.strip())
            pdf.close()
        except Exception:
            return ""
        return "\n".join(t for t in texts if t)

    def search(self, query, collection_id=None, top_k=5):
        q_emb = list(self.embedder.embed([query]))
        if not q_emb:
            return []
        import numpy as np
        q_vec = np.array(q_emb[0], dtype=np.float32)

        db = self._db()
        if collection_id:
            rows = db.execute("""
                SELECT rc.content, rc.document_id, rd.filename, rd.collection_id, rc.embedding
                FROM rag_chunks rc
                JOIN rag_documents rd ON rc.document_id = rd.id
                WHERE rd.collection_id = ?
            """, (collection_id,)).fetchall()
        else:
            rows = db.execute("""
                SELECT rc.content, rc.document_id, rd.filename, rd.collection_id, rc.embedding
                FROM rag_chunks rc
                JOIN rag_documents rd ON rc.document_id = rd.id
            """).fetchall()
        db.close()

        results = []
        for row in rows:
            if row[4] is None:
                continue
            chunk_vec = np.frombuffer(row[4], dtype=np.float32)
            if len(chunk_vec) != len(q_vec):
                continue
            similarity = float(np.dot(q_vec, chunk_vec) / (np.linalg.norm(q_vec) * np.linalg.norm(chunk_vec)))
            results.append({
                "content": row[0][:500],
                "document_id": row[1], "filename": row[2],
                "collection_id": row[3], "score": round(similarity * 100)
            })
        results.sort(key=lambda x: x["score"], reverse=True)
        return results[:top_k]

    def list_documents(self, collection_id=None):
        db = self._db()
        if collection_id:
            rows = db.execute(
                "SELECT id,filename,mime_type,char_count,chunk_count,content_preview,created_at FROM rag_documents WHERE collection_id=? ORDER BY created_at DESC",
                (collection_id,)
            ).fetchall()
        else:
            rows = db.execute(
                "SELECT id,filename,mime_type,char_count,chunk_count,content_preview,created_at FROM rag_documents ORDER BY created_at DESC"
            ).fetchall()
        db.close()
        return [{"id": r[0], "filename": r[1], "mime_type": r[2], "char_count": r[3],
                 "chunk_count": r[4], "preview": r[5], "created_at": r[6]} for r in rows]

    def delete_document(self, doc_id):
        db = self._db()
        db.execute("DELETE FROM rag_chunks WHERE document_id=?", (doc_id,))
        db.execute("DELETE FROM rag_documents WHERE id=?", (doc_id,))
        db.commit()
        db.close()

    def stats(self):
        db = self._db()
        docs = db.execute("SELECT COUNT(*) FROM rag_documents").fetchone()[0]
        chunks = db.execute("SELECT COUNT(*) FROM rag_chunks").fetchone()[0]
        colls = db.execute("SELECT COUNT(*) FROM rag_collections").fetchone()[0]
        db.close()
        return {"documents": docs, "chunks": chunks, "collections": colls}
